#!/usr/bin/env python3
"""Generate the SID Architecture Blueprint document as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color palette (SID brand) ──
GOLD = RGBColor(0xD4, 0xA9, 0x6A)       # SID gold
PURPLE = RGBColor(0x9B, 0x7E, 0xD8)     # SID purple
DARK = RGBColor(0x1A, 0x1A, 0x2E)       # Dark navy
MEDIUM = RGBColor(0x3A, 0x3A, 0x5E)     # Medium navy
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xFA)
TEXT_PRIMARY = RGBColor(0x2D, 0x2D, 0x3F)
TEXT_SECONDARY = RGBColor(0x6B, 0x6B, 0x8A)
ACCENT_GREEN = RGBColor(0x4E, 0xC9, 0xA3)
ACCENT_RED = RGBColor(0xE8, 0x6B, 0x6B)
ACCENT_BLUE = RGBColor(0x5B, 0xA3, 0xE8)

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = TEXT_PRIMARY

# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = DARK
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = PURPLE
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = GOLD
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, spacing_after=6, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(spacing_after)
    if alignment:
        p.alignment = alignment
    return p

def add_table_with_style(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="3A3A5E"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = "F5F5FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = TEXT_PRIMARY
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{bg}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacer
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="D4A96A"/></w:pBdr>')
    # Insert pBdr before any rPr or other late elements
    pPr.insert(0, pBdr)


# ════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SID')
run.font.size = Pt(48)
run.font.color.rgb = GOLD
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ARCHITECTURE BLUEPRINT')
run.font.size = Pt(18)
run.font.color.rgb = PURPLE
run.font.name = 'Calibri'
run.bold = True
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Vendor Management Platform')
run.font.size = Pt(13)
run.font.color.rgb = TEXT_SECONDARY
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(40)

add_divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('Version 2.0 \u2014 Complete System Redesign')
run.font.size = Pt(11)
run.font.color.rgb = TEXT_SECONDARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('March 2026')
run.font.size = Pt(11)
run.font.color.rgb = TEXT_SECONDARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('Prepared for Melissa Allen')
run.font.size = Pt(11)
run.font.color.rgb = GOLD
run.italic = True

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════

add_heading_styled('Table of Contents', level=1)
toc_items = [
    ('1', 'Executive Summary'),
    ('2', 'Current System Analysis'),
    ('3', 'Target Architecture'),
    ('4', 'Data Model Design'),
    ('5', 'Manager Dashboard (Your Side)'),
    ('6', 'Vendor Portal (Their Side)'),
    ('7', 'The SID Engine \u2014 Intelligence Layer'),
    ('8', 'Authentication & Security'),
    ('9', 'Tech Stack Recommendation'),
    ('10', 'Deployment & Infrastructure'),
    ('11', 'Phased Build Plan'),
    ('12', 'Appendix: Current API Reference'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}.')
    run.font.color.rgb = GOLD
    run.font.size = Pt(11)
    run.bold = True
    run2 = p.add_run(f'  {title}')
    run2.font.size = Pt(11)
    run2.font.color.rgb = TEXT_PRIMARY
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════

add_heading_styled('1. Executive Summary', level=1)

add_para(
    'SID (Smart Intelligent Dashboard) is a vendor management platform built to streamline '
    'how Melissa manages her network of recruitment posting vendors. Today, SID handles the full '
    'lifecycle: assigning job postings to vendors, tracking what gets posted, ingesting Indeed/LinkedIn '
    'CSV exports to reconcile candidates, analyzing cost-per-applicant, and pushing real-time updates '
    'between the manager dashboard and individual vendor portals.',
    size=10.5
)

add_para(
    'The current system works \u2014 but it was built incrementally. A single 3,360-line Python server '
    'handles everything: API routing, file I/O, CSV parsing, fuzzy matching, AI generation, analytics, '
    'and static file serving. Data lives in JSON files on disk. There is no authentication, no database, '
    'no separation between frontend and backend, and no way to deploy it beyond localhost.',
    size=10.5
)

add_para(
    'This document lays out the complete blueprint for SID v2 \u2014 a production-grade rebuild that '
    'preserves every feature of the current system while solving its structural limitations. The goal: '
    'complex and advanced behind the scenes, but simple and beautiful for everyone who uses it.',
    size=10.5, bold=True
)

add_divider()

add_heading_styled('What This Blueprint Covers', level=3)
add_para('\u2022  Complete analysis of the current system \u2014 what works, what breaks, what\u2019s fragile', size=10)
add_para('\u2022  Target architecture \u2014 React frontend, Supabase backend, proper API layer', size=10)
add_para('\u2022  Data model design \u2014 every table, relationship, and field mapped out', size=10)
add_para('\u2022  Manager workflow \u2014 your complete daily process, automated', size=10)
add_para('\u2022  Vendor workflow \u2014 their complete experience, from login to CSV upload', size=10)
add_para('\u2022  The SID Engine \u2014 intelligence layer for matching, analytics, and AI generation', size=10)
add_para('\u2022  Authentication \u2014 PIN-based vendor login, secure manager access', size=10)
add_para('\u2022  Deployment plan \u2014 from localhost to production with real URLs', size=10)
add_para('\u2022  Phased build plan \u2014 5 phases, each delivering usable functionality', size=10)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 2. CURRENT SYSTEM ANALYSIS
# ════════════════════════════════════════════════════════════════════

add_heading_styled('2. Current System Analysis', level=1)

add_heading_styled('2.1 Architecture Overview', level=2)
add_para(
    'The current SID system is a monolithic Python HTTP server (sid_manager_server.py, ~3,360 lines) '
    'that serves both the API and all static HTML files. It uses Python\u2019s built-in http.server module '
    'with threading support \u2014 no frameworks like Flask or Django. All data is stored in JSON files '
    'organized by vendor name in subdirectories.',
    size=10.5
)

add_heading_styled('Current Stack', level=3)
add_table_with_style(
    ['Component', 'Technology', 'Assessment'],
    [
        ['Backend Server', 'Python http.server (stdlib)', 'Functional but fragile \u2014 no middleware, no validation framework'],
        ['Frontend (Manager)', 'Vanilla HTML/CSS/JS', 'Single monolithic HTML file (~5000+ lines)'],
        ['Frontend (Vendor Portal)', 'Vanilla HTML/CSS/JS', 'Template-generated per vendor'],
        ['Database', 'JSON files on disk', 'No indexing, no transactions, no concurrent safety'],
        ['Authentication', 'None', 'Anyone with the URL can access anything'],
        ['File Storage', 'Local filesystem', 'Uploads stored in ./uploads/{vendor}/'],
        ['Deployment', 'localhost:8888', 'Only accessible on Melissa\u2019s machine'],
        ['AI Integration', 'Anthropic/OpenAI API', 'Direct HTTP calls for job description generation'],
        ['CSV Processing', 'Custom Python parser', 'Robust but tightly coupled to server'],
    ],
    col_widths=[1.8, 2.0, 2.7]
)

add_heading_styled('2.2 What Works Well', level=2)
add_para('\u2713 Complete posting lifecycle \u2014 assign, post, track, close, analyze', size=10, color=ACCENT_GREEN)
add_para('\u2713 Intelligent CSV reconciliation with 4-tier matching (ID \u2192 exact \u2192 fuzzy \u2192 unmatched)', size=10, color=ACCENT_GREEN)
add_para('\u2713 Posting Runs abstraction \u2014 links assignments to CSV data to candidates to analytics', size=10, color=ACCENT_GREEN)
add_para('\u2713 Multi-currency budget tracking with automatic USD normalization', size=10, color=ACCENT_GREEN)
add_para('\u2713 Alert system for pushing real-time updates to vendor portals', size=10, color=ACCENT_GREEN)
add_para('\u2713 Bulletin board with date-based history for daily instructions', size=10, color=ACCENT_GREEN)
add_para('\u2713 Auto-alert generation when manager modifies assignment fields', size=10, color=ACCENT_GREEN)
add_para('\u2713 Comprehensive analytics aggregation (by posting, category, account, daily)', size=10, color=ACCENT_GREEN)
add_para('\u2713 AI-powered job description generation with configurable API keys', size=10, color=ACCENT_GREEN)

add_heading_styled('2.3 Critical Weaknesses', level=2)
add_para('\u2717 No authentication \u2014 any vendor can see any other vendor\u2019s data', size=10, color=ACCENT_RED)
add_para('\u2717 JSON file storage \u2014 no concurrent write safety, no indexing, no querying', size=10, color=ACCENT_RED)
add_para('\u2717 Single monolithic server \u2014 one crash takes down everything', size=10, color=ACCENT_RED)
add_para('\u2717 Localhost only \u2014 vendors cannot access their portals remotely', size=10, color=ACCENT_RED)
add_para('\u2717 Manual multipart parsing \u2014 fragile file upload handling', size=10, color=ACCENT_RED)
add_para('\u2717 No real-time updates \u2014 vendors must refresh to see changes', size=10, color=ACCENT_RED)
add_para('\u2717 No backup/recovery \u2014 a corrupted JSON file means lost data', size=10, color=ACCENT_RED)
add_para('\u2717 Hardcoded vendor references \u2014 some paths assume specific vendor names', size=10, color=ACCENT_RED)
add_para('\u2717 No audit trail \u2014 no way to track who changed what and when', size=10, color=ACCENT_RED)

add_heading_styled('2.4 Current API Surface', level=2)
add_para(
    'The server exposes approximately 30 API endpoints across GET, POST, PATCH, and DELETE methods. '
    'This is actually a well-thought-out API design that maps cleanly to the business domain. '
    'The endpoint structure will carry forward into v2 with minimal changes.',
    size=10.5
)

add_table_with_style(
    ['Domain', 'Endpoints', 'Purpose'],
    [
        ['Vendors', '4', 'List, create, get accounts, save accounts'],
        ['Postings', '2', 'Get/save posting grid per vendor'],
        ['Assignments', '5', 'Create, read, update, delete, list all'],
        ['Activity', '4', 'Post/close tracking, undo, list all'],
        ['Alerts', '2', 'Push alerts to vendors, dismiss'],
        ['Bulletin', '2', 'Daily task/note boards per vendor'],
        ['Runs', '3', 'Posting lifecycle tracking, confirm links, title mapping'],
        ['Candidates', '2', 'Applicant data from CSV imports'],
        ['Analytics', '2', 'Per-vendor and cross-vendor analytics'],
        ['Uploads', '3', 'File upload, peek preview, process CSV'],
        ['AI Config', '2', 'Save/load AI generation settings'],
        ['Misc', '3', 'Announcements, git push, static files'],
    ],
    col_widths=[1.5, 1.0, 4.0]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 3. TARGET ARCHITECTURE
# ════════════════════════════════════════════════════════════════════

add_heading_styled('3. Target Architecture', level=1)

add_para(
    'SID v2 moves from a monolithic Python server to a modern three-tier architecture. '
    'The frontend becomes a React application with separate builds for the manager dashboard '
    'and vendor portal. The backend becomes a proper API layer. The database moves from JSON files '
    'to Supabase (hosted PostgreSQL with built-in auth, storage, and real-time subscriptions).',
    size=10.5
)

add_heading_styled('3.1 High-Level Architecture', level=2)

# Architecture diagram as a table
add_para('SYSTEM ARCHITECTURE', bold=True, size=11, color=PURPLE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')

add_table_with_style(
    ['Layer', 'Component', 'Technology', 'Responsibility'],
    [
        ['Presentation', 'Manager Dashboard', 'React + Tailwind CSS', 'Melissa\u2019s control center \u2014 assign, track, analyze'],
        ['Presentation', 'Vendor Portal', 'React + Tailwind CSS', 'Vendor\u2019s workspace \u2014 view assignments, upload CSVs, see bulletins'],
        ['Presentation', 'SID Mascot UI', 'React component', 'Animated SID with eye tracking, accessories, personality'],
        ['API Layer', 'REST API', 'Supabase Edge Functions or Next.js API Routes', 'Business logic, validation, authorization'],
        ['Intelligence', 'SID Engine', 'TypeScript modules', 'CSV parsing, fuzzy matching, analytics aggregation, AI generation'],
        ['Data', 'PostgreSQL', 'Supabase (hosted)', 'All structured data \u2014 vendors, assignments, activity, candidates'],
        ['Data', 'File Storage', 'Supabase Storage', 'CSV uploads, processed files, vendor assets'],
        ['Auth', 'Authentication', 'Supabase Auth + custom PINs', 'Manager login, vendor PIN access, row-level security'],
        ['Realtime', 'Subscriptions', 'Supabase Realtime', 'Live updates: assignment changes push to vendor portals instantly'],
        ['Deployment', 'Hosting', 'Vercel (frontend) + Supabase (backend)', 'Production URLs, SSL, CDN, auto-scaling'],
    ],
    col_widths=[1.1, 1.5, 2.0, 1.9]
)

add_heading_styled('3.2 Why This Stack', level=2)

add_para('Supabase as the backend:', bold=True, size=10.5)
add_para(
    'Supabase gives us PostgreSQL (real database with indexing, transactions, and querying), '
    'built-in authentication, file storage with access control, real-time subscriptions via WebSockets, '
    'and row-level security policies \u2014 all without writing a custom backend server. This eliminates '
    'the biggest weakness of the current system (JSON files, no auth, no real-time) in one move.',
    size=10
)

add_para('React + Tailwind for the frontend:', bold=True, size=10.5)
add_para(
    'React gives us component-based architecture, proper state management, and the ability to build '
    'the manager dashboard and vendor portal as separate apps sharing a common component library. '
    'Tailwind CSS provides utility-first styling that\u2019s fast to build with and easy to keep consistent. '
    'The SID mascot, calendar, assignment cards, and analytics charts all become reusable components.',
    size=10
)

add_para('Vercel for deployment:', bold=True, size=10.5)
add_para(
    'Vercel deploys React apps with zero configuration. Each vendor gets a clean URL like '
    'sid.yourdomain.com/portal/vendorname. Automatic SSL, CDN caching for static assets, '
    'and preview deployments for testing changes before they go live.',
    size=10
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 4. DATA MODEL DESIGN
# ════════════════════════════════════════════════════════════════════

add_heading_styled('4. Data Model Design', level=1)

add_para(
    'The data model translates every JSON file in the current system into proper database tables '
    'with relationships, constraints, and indexes. Every field from the current system is preserved.',
    size=10.5
)

add_heading_styled('4.1 Core Tables', level=2)

# Vendors table
add_para('vendors', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated unique identifier'],
        ['folder_name', 'TEXT (unique)', 'URL-safe name, e.g. "FBSPL"'],
        ['display_name', 'TEXT', 'Human-readable name, e.g. "FlatBridge Solutions"'],
        ['status', 'ENUM', 'active, paused, archived'],
        ['timezone', 'TEXT', 'Vendor\u2019s timezone, e.g. "Asia/Kolkata"'],
        ['pin_hash', 'TEXT', 'Hashed 6-digit portal access PIN'],
        ['onboarding_complete', 'BOOLEAN', 'Whether vendor has finished setup checklist'],
        ['portal_theme', 'JSONB', 'Custom portal appearance settings'],
        ['created_at', 'TIMESTAMPTZ', 'When vendor was added'],
        ['updated_at', 'TIMESTAMPTZ', 'Last modification timestamp'],
    ],
    col_widths=[1.8, 1.5, 3.2]
)

# Vendor Accounts table
add_para('vendor_accounts', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['vendor_id', 'UUID (FK \u2192 vendors)', 'Which vendor owns this account'],
        ['platform', 'TEXT', '"Indeed", "LinkedIn", etc.'],
        ['label', 'TEXT', 'Display label, e.g. "Indeed (USD)"'],
        ['account_id_external', 'TEXT', 'External account identifier'],
        ['currency', 'TEXT', 'USD, INR, EUR, etc.'],
        ['is_active', 'BOOLEAN', 'Whether this account is currently in use'],
    ],
    col_widths=[1.8, 1.8, 2.9]
)

# Postings table
add_para('postings', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['vendor_id', 'UUID (FK \u2192 vendors)', 'Assigned vendor'],
        ['grid_id', 'TEXT', 'Original grid ID, e.g. "SFG-0"'],
        ['category', 'TEXT', 'Posting category, e.g. "SFG"'],
        ['title', 'TEXT', 'Job title'],
        ['description', 'TEXT', 'Full job description (HTML)'],
        ['location', 'JSONB', '{type, city, state, remote, hybrid}'],
        ['salary', 'TEXT', 'Salary range string'],
        ['benefits', 'JSONB', 'Array of benefit strings'],
        ['budget', 'TEXT', 'Budget instructions, e.g. "$5/day, $10 max"'],
        ['account_id', 'UUID (FK \u2192 vendor_accounts)', 'Which platform account to post on'],
        ['status', 'ENUM', 'draft, active, posted, closed, cancelled, paused'],
        ['close_date', 'DATE', 'Target close date'],
        ['assigned_at', 'TIMESTAMPTZ', 'When assignment was pushed'],
        ['posted_at', 'TIMESTAMPTZ', 'When vendor marked as posted'],
        ['closed_at', 'TIMESTAMPTZ', 'When posting was closed'],
        ['created_at', 'TIMESTAMPTZ', 'Record creation time'],
        ['updated_at', 'TIMESTAMPTZ', 'Last modification'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

# Activity table
add_para('activity_log', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['posting_id', 'UUID (FK \u2192 postings)', 'Which posting this activity is for'],
        ['vendor_id', 'UUID (FK \u2192 vendors)', 'Which vendor'],
        ['action', 'ENUM', 'assigned, posted, closed, budget_change, note_added, status_change'],
        ['actor', 'TEXT', '"manager" or "vendor"'],
        ['details', 'JSONB', 'Action-specific data (old/new values, links, etc.)'],
        ['created_at', 'TIMESTAMPTZ', 'When this activity occurred'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_page_break()

# Candidates table
add_para('candidates', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['posting_id', 'UUID (FK \u2192 postings)', 'Linked posting (may be null if unmatched)'],
        ['vendor_id', 'UUID (FK \u2192 vendors)', 'Which vendor uploaded this data'],
        ['csv_upload_id', 'UUID (FK \u2192 csv_uploads)', 'Source CSV file'],
        ['name', 'TEXT', 'Candidate full name'],
        ['email', 'TEXT', 'Candidate email'],
        ['phone', 'TEXT', 'Candidate phone'],
        ['job_title_csv', 'TEXT', 'Job title as it appeared in the CSV'],
        ['location', 'TEXT', 'Candidate location string'],
        ['is_us_based', 'BOOLEAN', 'Whether location resolves to US'],
        ['source', 'TEXT', 'Sponsored, Organic, etc.'],
        ['applied_at', 'TIMESTAMPTZ', 'Application date from CSV'],
        ['resume_url', 'TEXT', 'Link to resume if available'],
        ['raw_data', 'JSONB', 'Complete original CSV row for reference'],
        ['created_at', 'TIMESTAMPTZ', 'When imported into SID'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

# CSV Uploads table
add_para('csv_uploads', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['vendor_id', 'UUID (FK \u2192 vendors)', 'Uploading vendor'],
        ['filename', 'TEXT', 'Original filename'],
        ['storage_path', 'TEXT', 'Path in Supabase Storage'],
        ['file_size_bytes', 'INTEGER', 'File size'],
        ['row_count', 'INTEGER', 'Number of data rows detected'],
        ['status', 'ENUM', 'uploaded, processing, processed, error'],
        ['matched_count', 'INTEGER', 'Rows successfully matched to postings'],
        ['unmatched_count', 'INTEGER', 'Rows that couldn\u2019t be matched'],
        ['processing_log', 'JSONB', 'Detailed processing results and errors'],
        ['uploaded_at', 'TIMESTAMPTZ', 'When the file was uploaded'],
        ['processed_at', 'TIMESTAMPTZ', 'When processing completed'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

# Posting Runs table
add_para('posting_runs', bold=True, size=11, color=PURPLE)
add_para(
    'This is the core intelligence table. A "run" represents one complete lifecycle of a posting: '
    'from assignment through posting, candidate collection, and closing. It links assignments to CSV '
    'data to candidates to analytics, enabling cost-per-applicant tracking and performance analysis.',
    size=10, italic=True
)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['posting_id', 'UUID (FK \u2192 postings)', 'The posting this run tracks'],
        ['vendor_id', 'UUID (FK \u2192 vendors)', 'Vendor'],
        ['run_number', 'INTEGER', 'Sequential run number for this posting'],
        ['status', 'ENUM', 'active, posted, closed'],
        ['candidates_total', 'INTEGER', 'Total candidates linked to this run'],
        ['candidates_us', 'INTEGER', 'US-based candidates'],
        ['candidates_intl', 'INTEGER', 'International candidates'],
        ['budget_spent_usd', 'DECIMAL', 'Total spend normalized to USD'],
        ['budget_spent_raw', 'TEXT', 'Original currency amount string'],
        ['cost_per_applicant', 'DECIMAL', 'Calculated CPA in USD'],
        ['sources', 'JSONB', '{"Sponsored": 45, "Organic": 42}'],
        ['days_active', 'INTEGER', 'Number of days this run was live'],
        ['csv_job_titles', 'TEXT[]', 'CSV titles matched to this run'],
        ['closed_at', 'TIMESTAMPTZ', 'When this run was closed'],
        ['created_at', 'TIMESTAMPTZ', 'When run was created'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

# Additional tables
add_heading_styled('4.2 Supporting Tables', level=2)

add_para('alerts', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['vendor_id', 'UUID (FK)', 'Target vendor'],
        ['type', 'ENUM', 'missing_csv, action_required, info, warning'],
        ['priority', 'ENUM', 'low, normal, high, urgent'],
        ['title', 'TEXT', 'Alert headline'],
        ['message', 'TEXT', 'Alert body text'],
        ['metadata', 'JSONB', 'Posting references, update details, etc.'],
        ['is_dismissed', 'BOOLEAN', 'Whether vendor has dismissed this alert'],
        ['created_at', 'TIMESTAMPTZ', 'When alert was pushed'],
        ['dismissed_at', 'TIMESTAMPTZ', 'When dismissed'],
    ],
    col_widths=[1.5, 1.8, 3.2]
)

add_para('bulletins', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['vendor_id', 'UUID (FK)', 'Target vendor'],
        ['date', 'DATE', 'Bulletin date'],
        ['tasks', 'JSONB', 'Array of task strings'],
        ['notes', 'TEXT', 'Free-form notes for the day'],
        ['created_by', 'TEXT', '"manager" or "system"'],
        ['created_at', 'TIMESTAMPTZ', 'When created'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

add_para('announcements', bold=True, size=11, color=PURPLE)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['title', 'TEXT', 'Announcement title'],
        ['body', 'TEXT', 'Announcement content'],
        ['target', 'ENUM', 'all_vendors, specific_vendor, managers_only'],
        ['vendor_ids', 'UUID[]', 'If target is specific_vendor'],
        ['is_active', 'BOOLEAN', 'Whether currently visible'],
        ['expires_at', 'TIMESTAMPTZ', 'Auto-expire date (optional)'],
        ['created_at', 'TIMESTAMPTZ', 'When created'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

add_para('title_mappings', bold=True, size=11, color=PURPLE)
add_para(
    'Manual overrides for when fuzzy matching can\u2019t resolve a CSV job title to a posting. '
    'Once a mapping is confirmed, it persists and auto-applies to future CSV imports.',
    size=10, italic=True
)
add_table_with_style(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'UUID (PK)', 'Auto-generated'],
        ['vendor_id', 'UUID (FK)', 'Vendor this mapping applies to'],
        ['csv_job_title', 'TEXT', 'The job title as it appears in CSVs'],
        ['posting_id', 'UUID (FK \u2192 postings)', 'The posting it maps to'],
        ['created_by', 'TEXT', 'Who created this mapping'],
        ['created_at', 'TIMESTAMPTZ', 'When mapping was established'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 5. MANAGER DASHBOARD
# ════════════════════════════════════════════════════════════════════

add_heading_styled('5. Manager Dashboard (Your Side)', level=1)

add_para(
    'The manager dashboard is your command center. Everything you do today \u2014 assigning postings, '
    'tracking activity, reviewing CSV data, pushing updates to vendors \u2014 lives here. The v2 rebuild '
    'turns the current monolithic HTML file into a proper React application with distinct views.',
    size=10.5
)

add_heading_styled('5.1 Dashboard Home', level=2)
add_para('What you see when you open SID:', bold=True, size=10.5)
add_para('\u2022  SID mascot in the sidebar corner with eye tracking and seasonal accessories', size=10)
add_para('\u2022  Personalized greeting with your name and time-appropriate message', size=10)
add_para('\u2022  Dual timezone clock (Eastern Time primary, vendor\u2019s local time below)', size=10)
add_para('\u2022  Week strip calendar with color-coded activity dots per day', size=10)
add_para('\u2022  Today\u2019s key stats: Active Assignments, Posted Today, Close Soon, Pending CSVs', size=10)
add_para('\u2022  SID thought bubble with rotating personality-driven tips', size=10)
add_para('\u2022  VENDOR HUB branding in sidebar footer with dark mode toggle', size=10)

add_heading_styled('5.2 Core Views', level=2)

add_table_with_style(
    ['View', 'Purpose', 'Key Features'],
    [
        ['Posting Grid', 'Master list of all postings per vendor', 'Category grouping, drag-to-assign, bulk actions, AI description generator'],
        ['Assignment Manager', 'Track what\u2019s been pushed to vendors', 'Status pipeline (assigned \u2192 posted \u2192 closed), date filters, auto-alerts on changes'],
        ['CSV Reconciliation', 'Match uploaded CSVs to postings', '4-tier matching engine, manual override UI, visual match confidence scores'],
        ['Analytics', 'Performance data and spending', 'Cost-per-applicant charts, category breakdowns, budget tracking, trend lines'],
        ['Vendor Manager', 'Add/configure/manage vendors', 'Create vendor, set PIN, configure accounts, view onboarding status'],
        ['Bulletin Board', 'Daily instructions per vendor', 'Date-based history, task lists, free-form notes, template-based quick fill'],
        ['Run Tracker', 'Posting lifecycle dashboard', 'See every run: candidates linked, budget spent, CPA, days active'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

add_heading_styled('5.3 Daily Manager Workflow', level=2)
add_para(
    'This is your actual daily process, mapped to SID features. The v2 system automates or streamlines '
    'every step:',
    size=10.5
)

add_para('Morning:', bold=True, size=10.5, color=GOLD)
add_para('1. Open SID \u2192 Dashboard shows today\u2019s stats, SID greets you by name', size=10)
add_para('2. Check calendar \u2192 See which postings are due to close this week', size=10)
add_para('3. Review overnight CSV uploads \u2192 Auto-processed with match results waiting', size=10)
add_para('4. Resolve any unmatched titles \u2192 Title mapping UI with suggestions', size=10)

add_para('Midday:', bold=True, size=10.5, color=GOLD)
add_para('5. Check posting statuses \u2192 See which vendors have marked postings as posted', size=10)
add_para('6. Push new assignments \u2192 Select from grid, assign to vendor, auto-alert sent', size=10)
add_para('7. Update bulletins \u2192 Write daily instructions for each vendor', size=10)

add_para('Afternoon/Evening:', bold=True, size=10.5, color=GOLD)
add_para('8. Review analytics \u2192 Check CPA trends, identify underperforming postings', size=10)
add_para('9. Adjust budgets/close dates \u2192 Changes auto-alert vendors', size=10)
add_para('10. End-of-day review \u2192 Dashboard summary of today\u2019s completed activity', size=10)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 6. VENDOR PORTAL
# ════════════════════════════════════════════════════════════════════

add_heading_styled('6. Vendor Portal (Their Side)', level=1)

add_para(
    'The vendor portal is what your vendors see. It needs to be dead simple, clear about what they '
    'need to do, and impossible to break. Each vendor gets their own portal with only their data visible.',
    size=10.5
)

add_heading_styled('6.1 Vendor Login Flow', level=2)
add_para('1. Vendor navigates to sid.yourdomain.com/portal/vendorname', size=10)
add_para('2. Enters their 6-digit PIN (you set this when creating the vendor)', size=10)
add_para('3. PIN is verified against Supabase Auth \u2014 returns a session token', size=10)
add_para('4. Session persists for 7 days (configurable) \u2014 no re-login needed', size=10)
add_para('5. All API calls include the session token \u2014 row-level security ensures they only see their data', size=10)

add_heading_styled('6.2 Vendor Dashboard', level=2)
add_para('What vendors see when they log in:', bold=True, size=10.5)
add_para('\u2022  SID mascot greeting them by name with fun personality messages', size=10)
add_para('\u2022  Active alerts from you (changes to postings, missing CSVs, action required)', size=10)
add_para('\u2022  Today\u2019s bulletin with task list and notes from you', size=10)
add_para('\u2022  Assignment cards showing what they need to post today', size=10)
add_para('\u2022  Quick-action buttons: Upload CSV, View Guide, Mark Posted', size=10)
add_para('\u2022  Announcements banner for system-wide messages', size=10)
add_para('\u2022  Extension install prompt and onboarding checklist (new vendors)', size=10)

add_heading_styled('6.3 Vendor Core Features', level=2)
add_table_with_style(
    ['Feature', 'Description', 'Current Status \u2192 v2 Improvement'],
    [
        ['Assignment View', 'See all active assignments with full details', 'Works \u2192 Add real-time updates, status timeline'],
        ['CSV Upload', 'Upload Indeed/LinkedIn CSV exports', 'Works \u2192 Add drag-drop, auto-process, progress bar'],
        ['Mark as Posted', 'Confirm a posting is live with link', 'Works \u2192 Add link validation, auto-detect platform'],
        ['Bulletin Board', 'Read daily instructions from manager', 'Works \u2192 Add task checkoff, read receipts'],
        ['Alert Inbox', 'See and dismiss manager alerts', 'Works \u2192 Add priority sorting, action buttons'],
        ['Performance View', 'See their own posting performance', 'Partial \u2192 Full analytics with charts and trends'],
        ['Guide/SOP', 'Step-by-step posting instructions', 'Works \u2192 Interactive walkthrough with screenshots'],
    ],
    col_widths=[1.5, 2.2, 2.8]
)

add_heading_styled('6.4 Daily Vendor Workflow', level=2)
add_para('1. Open portal \u2192 See alerts first (anything urgent from Melissa)', size=10)
add_para('2. Check bulletin \u2192 Read today\u2019s instructions and task list', size=10)
add_para('3. View assignments \u2192 See what needs to be posted today', size=10)
add_para('4. Post jobs on Indeed/LinkedIn \u2192 Follow the guide for each posting', size=10)
add_para('5. Mark as posted \u2192 Click "Posted" with the live link \u2014 Melissa sees it instantly', size=10)
add_para('6. Upload CSV \u2192 Drag-drop the Indeed/LinkedIn export \u2014 auto-processed', size=10)
add_para('7. Check for updates \u2192 Real-time: budget changes, close date changes appear live', size=10)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 7. THE SID ENGINE
# ════════════════════════════════════════════════════════════════════

add_heading_styled('7. The SID Engine \u2014 Intelligence Layer', level=1)

add_para(
    'The SID Engine is the brain behind the dashboard. It handles everything that requires intelligence: '
    'matching CSV data to postings, aggregating analytics, generating job descriptions, and detecting '
    'anomalies. In v1, this logic is scattered across 15+ server methods. In v2, it becomes a clean '
    'set of TypeScript modules.',
    size=10.5
)

add_heading_styled('7.1 CSV Reconciliation Engine', level=2)
add_para(
    'This is the most complex piece of SID. When a vendor uploads an Indeed CSV, the engine must '
    'figure out which posting each row belongs to. The current 4-tier matching system is excellent '
    'and carries forward:',
    size=10.5
)

add_para('Tier 0 \u2014 Assignment ID Match:', bold=True, size=10.5, color=GOLD)
add_para(
    'If the CSV contains an Assignment_ID column (from SID\u2019s own CSV export), it\u2019s an instant, '
    'guaranteed match. No ambiguity.',
    size=10
)

add_para('Tier 1 \u2014 Manual Title Map:', bold=True, size=10.5, color=GOLD)
add_para(
    'If a previous CSV had an unmatched title and the manager manually mapped it, that mapping '
    'is stored in title_mappings and auto-applied to all future CSVs.',
    size=10
)

add_para('Tier 2 \u2014 Exact Title Match:', bold=True, size=10.5, color=GOLD)
add_para(
    'The CSV job title matches a posting title exactly (case-insensitive). Most CSVs match here.',
    size=10
)

add_para('Tier 3 \u2014 Fuzzy Title Match:', bold=True, size=10.5, color=GOLD)
add_para(
    'Jaccard similarity on tokenized, lowercased words with a substring bonus. Threshold: 0.6. '
    'Example: "Insurance Agent - Remote (Work From Home)" fuzzy-matches to "Insurance Agent" '
    'at ~0.85 confidence.',
    size=10
)

add_para('Tier 4 \u2014 Unmatched:', bold=True, size=10.5, color=GOLD)
add_para(
    'If nothing matches above 0.6, the title goes to the unmatched queue. The manager sees it '
    'in the reconciliation UI and can manually map it (which persists for future imports).',
    size=10
)

add_heading_styled('7.2 Analytics Engine', level=2)
add_para(
    'The analytics engine aggregates data across all posting runs to produce actionable insights. '
    'It currently handles: daily activity summaries, per-posting performance (CPA, applicant count, spend), '
    'per-category breakdowns, per-account spending, and multi-currency normalization to USD. '
    'In v2, this moves to SQL views and materialized queries for instant loading.',
    size=10.5
)

add_heading_styled('7.3 AI Generation Engine', level=2)
add_para(
    'The current system supports both Anthropic (Claude) and OpenAI for generating job descriptions. '
    'In v2, this extends to: auto-generating bulletin tasks based on today\u2019s assignments, suggesting '
    'budget adjustments based on CPA trends, drafting close-date recommendations, and '
    'anomaly detection (e.g., vendor hasn\u2019t uploaded a CSV in 3 days).',
    size=10.5
)

add_heading_styled('7.4 Applicant Normalization', level=2)
add_para(
    'The current system already has sophisticated applicant row normalization (~100 lines) that handles '
    'different Indeed CSV formats, extracts location/source/date fields, and detects US vs. international '
    'candidates. This logic carries forward with additional support for LinkedIn CSV format variants.',
    size=10.5
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 8. AUTHENTICATION & SECURITY
# ════════════════════════════════════════════════════════════════════

add_heading_styled('8. Authentication & Security', level=1)

add_para(
    'The biggest gap in the current system is zero authentication. Anyone with the URL can see everything. '
    'V2 fixes this completely with two separate auth flows.',
    size=10.5
)

add_heading_styled('8.1 Manager Authentication', level=2)
add_para(
    'You (and any future team members) log in with email + password via Supabase Auth. This gives you '
    'full access to all vendors, all data, all admin functions. Supabase Auth handles session management, '
    'token refresh, and secure cookie storage.',
    size=10.5
)

add_heading_styled('8.2 Vendor PIN Authentication', level=2)
add_para(
    'Vendors get a simpler flow because they\u2019re not tech-savvy users and we don\u2019t want password reset '
    'friction. Each vendor gets a 6-digit PIN that you set when creating them. The flow:',
    size=10.5
)
add_para('1. Vendor navigates to their portal URL', size=10)
add_para('2. Enters 6-digit PIN on a clean login screen with SID mascot', size=10)
add_para('3. PIN is verified server-side (bcrypt hashed, never stored in plain text)', size=10)
add_para('4. On success, a Supabase session is created scoped to that vendor\u2019s data', size=10)
add_para('5. Session persists for 7 days \u2014 vendor doesn\u2019t need to re-enter PIN daily', size=10)
add_para('6. You can reset/change any vendor\u2019s PIN from the manager dashboard', size=10)

add_heading_styled('8.3 Row-Level Security', level=2)
add_para(
    'Supabase\u2019s row-level security (RLS) ensures that even if someone somehow gets a valid session token, '
    'they can only see data for their own vendor. Every table has RLS policies:',
    size=10.5
)
add_para('\u2022  Vendors can only SELECT rows where vendor_id matches their session', size=10)
add_para('\u2022  Vendors can only INSERT into csv_uploads and activity_log for their vendor', size=10)
add_para('\u2022  Vendors cannot UPDATE or DELETE any data \u2014 only the manager can', size=10)
add_para('\u2022  Managers (identified by role in their JWT) bypass RLS and see everything', size=10)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 9. TECH STACK
# ════════════════════════════════════════════════════════════════════

add_heading_styled('9. Tech Stack Recommendation', level=1)

add_table_with_style(
    ['Layer', 'Technology', 'Why'],
    [
        ['Frontend Framework', 'React 18+ with TypeScript', 'Component architecture, huge ecosystem, strong typing'],
        ['UI Styling', 'Tailwind CSS', 'Fast development, consistent design, easy dark mode'],
        ['UI Components', 'shadcn/ui', 'Beautiful, accessible, customizable component library'],
        ['State Management', 'React Query (TanStack)', 'Server state caching, auto-refresh, optimistic updates'],
        ['Charts/Viz', 'Recharts', 'React-native charting, great for analytics dashboards'],
        ['Database', 'Supabase (PostgreSQL)', 'Managed database with auth, storage, real-time built in'],
        ['Authentication', 'Supabase Auth', 'Email/password for managers, custom PIN flow for vendors'],
        ['File Storage', 'Supabase Storage', 'CSV uploads with access control and signed URLs'],
        ['Real-time', 'Supabase Realtime', 'WebSocket subscriptions for live updates'],
        ['API Layer', 'Supabase Edge Functions + RPC', 'Server-side logic for CSV processing, AI calls'],
        ['AI Integration', 'Anthropic Claude API', 'Job description generation, smart suggestions'],
        ['Deployment', 'Vercel', 'Zero-config React hosting, preview deploys, custom domains'],
        ['Domain/SSL', 'Cloudflare', 'DNS management, SSL, DDoS protection, caching'],
        ['Monitoring', 'Supabase Dashboard + Vercel Analytics', 'Database performance, API latency, traffic'],
    ],
    col_widths=[1.5, 2.2, 2.8]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 10. DEPLOYMENT
# ════════════════════════════════════════════════════════════════════

add_heading_styled('10. Deployment & Infrastructure', level=1)

add_heading_styled('10.1 From Localhost to Production', level=2)
add_para(
    'The current system runs on localhost:8888. In production, the architecture deploys across '
    'two managed services that handle scaling, SSL, and availability automatically:',
    size=10.5
)

add_para('Supabase (Backend):', bold=True, size=10.5, color=PURPLE)
add_para('\u2022  PostgreSQL database with automatic backups', size=10)
add_para('\u2022  Edge Functions for server-side logic (CSV processing, AI calls)', size=10)
add_para('\u2022  Storage buckets for CSV files and vendor assets', size=10)
add_para('\u2022  Real-time WebSocket server for live updates', size=10)
add_para('\u2022  Auth service for session management', size=10)

add_para('Vercel (Frontend):', bold=True, size=10.5, color=PURPLE)
add_para('\u2022  Manager dashboard: sid.yourdomain.com', size=10)
add_para('\u2022  Vendor portals: sid.yourdomain.com/portal/{vendorname}', size=10)
add_para('\u2022  Automatic SSL certificates', size=10)
add_para('\u2022  CDN for static assets (React bundle, images, fonts)', size=10)
add_para('\u2022  Preview deployments for testing changes', size=10)

add_heading_styled('10.2 URL Structure', level=2)
add_table_with_style(
    ['URL', 'Who Sees It', 'What It Shows'],
    [
        ['sid.yourdomain.com', 'Melissa (Manager)', 'Manager dashboard with full control'],
        ['sid.yourdomain.com/portal/FBSPL', 'FBSPL vendor', 'FBSPL\u2019s portal with only their data'],
        ['sid.yourdomain.com/portal/vendorname', 'Each vendor', 'Their personalized portal'],
        ['sid.yourdomain.com/guide/vendorname', 'Each vendor', 'Their posting guide/SOP'],
        ['sid.yourdomain.com/api/*', 'Internal', 'API endpoints (proxied to Supabase)'],
    ],
    col_widths=[2.5, 1.5, 2.5]
)

add_heading_styled('10.3 Cost Estimate', level=2)
add_table_with_style(
    ['Service', 'Tier', 'Monthly Cost', 'What You Get'],
    [
        ['Supabase', 'Free tier (to start)', '$0', '500MB DB, 1GB storage, 50K auth users, 2M edge function calls'],
        ['Supabase', 'Pro (when needed)', '$25/mo', '8GB DB, 100GB storage, unlimited auth, 2M+ edge functions'],
        ['Vercel', 'Free tier (to start)', '$0', '100GB bandwidth, serverless functions, preview deploys'],
        ['Vercel', 'Pro (when needed)', '$20/mo', '1TB bandwidth, analytics, team features'],
        ['Cloudflare', 'Free tier', '$0', 'DNS, SSL, basic DDoS protection, caching'],
        ['Domain', 'Annual', '~$12/yr', 'Custom domain (e.g., sid-hub.com or vendorhub.app)'],
    ],
    col_widths=[1.2, 1.5, 1.0, 2.8]
)

add_para(
    'Starting cost: $0/month (free tiers cover everything for initial launch). '
    'Scale cost: ~$45/month when traffic and data grow beyond free tier limits.',
    size=10, italic=True, color=TEXT_SECONDARY
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 11. PHASED BUILD PLAN
# ════════════════════════════════════════════════════════════════════

add_heading_styled('11. Phased Build Plan', level=1)

add_para(
    'The rebuild happens in 5 phases. Each phase delivers working, usable functionality. '
    'You never have to wait until "everything is done" to start using new features.',
    size=10.5
)

add_divider()

add_heading_styled('Phase 1: Foundation', level=2)
add_para('Timeline: 1\u20132 weeks', bold=True, size=10, color=GOLD)
add_para('Goal: Set up the infrastructure and migrate data from JSON files to Supabase.', italic=True, size=10)
add_para('')
add_para('\u2022  Set up Supabase project with PostgreSQL database', size=10)
add_para('\u2022  Create all database tables with proper types, constraints, and indexes', size=10)
add_para('\u2022  Write data migration script: JSON files \u2192 Supabase tables', size=10)
add_para('\u2022  Set up Supabase Auth: manager email/password login', size=10)
add_para('\u2022  Set up Supabase Storage bucket for CSV uploads', size=10)
add_para('\u2022  Scaffold React project with Tailwind CSS and shadcn/ui', size=10)
add_para('\u2022  Deploy skeleton to Vercel with custom domain', size=10)
add_para('')
add_para('Deliverable: Infrastructure ready, data migrated, skeleton app deployed at your domain.', bold=True, size=10)

add_divider()

add_heading_styled('Phase 2: Manager Dashboard', level=2)
add_para('Timeline: 2\u20133 weeks', bold=True, size=10, color=GOLD)
add_para('Goal: Rebuild the manager dashboard in React with all current functionality.', italic=True, size=10)
add_para('')
add_para('\u2022  Dashboard home with SID mascot, greeting, calendar, stats', size=10)
add_para('\u2022  Posting grid with category grouping and assignment flow', size=10)
add_para('\u2022  Assignment manager with status pipeline view', size=10)
add_para('\u2022  Bulletin board editor per vendor', size=10)
add_para('\u2022  Vendor management (create, configure, set PINs)', size=10)
add_para('\u2022  Alert push system', size=10)
add_para('\u2022  Dark mode with SID brand theming', size=10)
add_para('')
add_para('Deliverable: You can manage all vendors from the new dashboard. Old system stays as backup.', bold=True, size=10)

add_divider()

add_heading_styled('Phase 3: Vendor Portal', level=2)
add_para('Timeline: 2\u20133 weeks', bold=True, size=10, color=GOLD)
add_para('Goal: Build vendor portals with PIN auth and real-time updates.', italic=True, size=10)
add_para('')
add_para('\u2022  PIN login screen with SID mascot', size=10)
add_para('\u2022  Vendor dashboard: alerts, bulletin, assignments, quick actions', size=10)
add_para('\u2022  CSV upload with drag-drop and auto-processing', size=10)
add_para('\u2022  "Mark as Posted" flow with link capture', size=10)
add_para('\u2022  Real-time updates via Supabase subscriptions', size=10)
add_para('\u2022  Row-level security: vendors only see their own data', size=10)
add_para('\u2022  Onboarding checklist for new vendors', size=10)
add_para('')
add_para('Deliverable: Vendors log in with PINs, see only their data, upload CSVs, and get live updates.', bold=True, size=10)

add_divider()

add_heading_styled('Phase 4: Intelligence Layer', level=2)
add_para('Timeline: 2\u20133 weeks', bold=True, size=10, color=GOLD)
add_para('Goal: Port the SID Engine logic and add new intelligence features.', italic=True, size=10)
add_para('')
add_para('\u2022  CSV reconciliation engine in TypeScript (all 4 tiers)', size=10)
add_para('\u2022  Title mapping UI with confidence scores', size=10)
add_para('\u2022  Analytics dashboard with charts (Recharts)', size=10)
add_para('\u2022  Posting Runs lifecycle tracking', size=10)
add_para('\u2022  Performance views: CPA trends, category breakdowns, budget tracking', size=10)
add_para('\u2022  AI job description generator via Edge Functions', size=10)
add_para('\u2022  Applicant normalization and US/international detection', size=10)
add_para('')
add_para('Deliverable: Full analytics, CSV reconciliation, and AI features working in production.', bold=True, size=10)

add_divider()

add_heading_styled('Phase 5: Polish & Advanced Features', level=2)
add_para('Timeline: 1\u20132 weeks', bold=True, size=10, color=GOLD)
add_para('Goal: Production hardening, advanced features, and the finishing touches.', italic=True, size=10)
add_para('')
add_para('\u2022  SID personality engine: seasonal accessories, expanded thought bubbles, animations', size=10)
add_para('\u2022  Team dashboard (if additional managers are added)', size=10)
add_para('\u2022  Email/Slack notifications for critical events', size=10)
add_para('\u2022  Automated anomaly detection (vendor hasn\u2019t uploaded CSV in X days)', size=10)
add_para('\u2022  Bulk operations: assign/close multiple postings at once', size=10)
add_para('\u2022  Export features: analytics reports, candidate lists, posting summaries', size=10)
add_para('\u2022  Performance optimization: lazy loading, caching, pagination', size=10)
add_para('\u2022  Comprehensive error handling and user-facing error messages', size=10)
add_para('')
add_para('Deliverable: Production-ready, polished, fully featured SID v2.', bold=True, size=10)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# 12. APPENDIX
# ════════════════════════════════════════════════════════════════════

add_heading_styled('12. Appendix: Current API Reference', level=1)

add_para(
    'Complete list of all API endpoints in the current sid_manager_server.py, organized by domain. '
    'These endpoints map directly to the v2 Supabase RPC functions and REST endpoints.',
    size=10.5
)

add_table_with_style(
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET', '/api/vendors', 'List all vendors (scans filesystem for vendor directories)'],
        ['POST', '/api/vendors', 'Create new vendor with full portal from template'],
        ['GET', '/api/vendor-accounts', 'Get platform account configuration per vendor'],
        ['POST', '/api/vendor-accounts', 'Save platform account configuration'],
        ['GET', '/api/postings/{vendor}', 'Get posting grid for a vendor'],
        ['POST', '/api/postings/{vendor}', 'Save posting grid for a vendor'],
        ['GET', '/api/assignments/{vendor}', 'Get assignments (today, by date, or date range)'],
        ['GET', '/api/assignments', 'Get all assignments across all vendors'],
        ['POST', '/api/assignments/{vendor}', 'Push new assignments to a vendor'],
        ['PATCH', '/api/assignments/{vendor}', 'Update assignment fields (auto-alerts vendor)'],
        ['DELETE', '/api/assignments/{vendor}', 'Cancel or hard-delete an assignment'],
        ['GET', '/api/activity/{vendor}', 'Get activity log for a vendor and date range'],
        ['GET', '/api/activity', 'Get all vendor activity'],
        ['POST', '/api/activity/{vendor}', 'Record post/close/undo actions'],
        ['GET', '/api/alerts/{vendor}', 'Get active alerts for a vendor'],
        ['POST', '/api/alerts/{vendor}', 'Create, dismiss, or dismiss-all alerts'],
        ['GET', '/api/bulletin/{vendor}', 'Get bulletin board with history'],
        ['POST', '/api/bulletin/{vendor}', 'Save bulletin for a vendor'],
        ['GET', '/api/announcements', 'Get system announcements'],
        ['POST', '/api/announcements', 'Save system announcements'],
        ['GET', '/api/runs/{vendor}', 'Get all posting runs for a vendor'],
        ['POST', '/api/runs/{vendor}/confirm', 'Confirm CSV-to-posting link mappings'],
        ['POST', '/api/runs/{vendor}/map', 'Save manual title mappings'],
        ['GET', '/api/candidates', 'Get all candidates (with filtering)'],
        ['GET', '/api/candidates/{vendor}', 'Get candidates for a specific vendor'],
        ['GET', '/api/performance/{vendor}', 'Get aggregated performance stats'],
        ['GET', '/api/analytics/{vendor}', 'Get full analytics data'],
        ['GET', '/api/analytics', 'Get analytics across all vendors'],
        ['POST', '/api/upload/{vendor}', 'Upload CSV file (multipart)'],
        ['GET', '/api/uploads/{vendor}', 'List uploaded files'],
        ['POST', '/api/peek-upload/{vendor}', 'Preview CSV contents before processing'],
        ['POST', '/api/process-upload/{vendor}', 'Process uploaded CSV with reconciliation'],
        ['GET', '/api/processed/{vendor}', 'Get processed CSV results'],
        ['GET', '/api/ai-config', 'Get AI generation configuration'],
        ['POST', '/api/ai-config', 'Save AI generation configuration'],
        ['POST', '/api/generate-posting', 'Generate job description via AI'],
        ['POST', '/api/push', 'Git add, commit, and push all changes'],
    ],
    col_widths=[0.8, 2.5, 3.2]
)

doc.add_page_break()

# ── Final page ──
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SID v2')
run.font.size = Pt(28)
run.font.color.rgb = GOLD
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Built to be complex behind the scenes.')
run.font.size = Pt(13)
run.font.color.rgb = TEXT_SECONDARY
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Simple and beautiful for everyone who uses it.')
run.font.size = Pt(13)
run.font.color.rgb = TEXT_SECONDARY
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('A well-oiled machine.')
run.font.size = Pt(14)
run.font.color.rgb = PURPLE
run.bold = True

# Fix zoom setting for validation
from lxml import etree
settings = doc.settings.element
zoom_elements = settings.findall(qn('w:zoom'))
for z in zoom_elements:
    if z.get(qn('w:percent')) is None:
        z.set(qn('w:percent'), '100')

# ── Save ──
output_path = '/sessions/friendly-inspiring-brahmagupta/mnt/SID/SID_Architecture_Blueprint.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
